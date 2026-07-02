import streamlit as st
import pandas as pd
import json
import os
import re
import base64
from datetime import datetime

FARMERS_DB = "farmers.json"

HS_CODE_MAP = {
    "Maize": "1005.90", "Coffee": "0901.11", "Tea": "0902.30",
    "Avocado": "0804.40", "French Beans": "0708.20", "Roses": "0603.11",
    "Macadamia Nuts": "0802.62", "Mango": "0804.50",
    "Pineapple": "0804.30", "Passion Fruit": "0810.90",
    "Snow Peas": "0710.21", "Carnations": "0603.12",
    "Spinach": "0709.70", "Kale": "0704.99",
    "Capsicum": "0709.60", "Tomato": "0702.00",
}

EUDR_RISK = {
    "Coffee": "AMBER", "Tea": "AMBER", "Maize": "AMBER",
    "Avocado": "GREEN", "Mango": "GREEN", "Macadamia Nuts": "GREEN",
    "French Beans": "GREEN", "Snow Peas": "GREEN", "Roses": "GREEN",
    "Carnations": "GREEN", "Passion Fruit": "GREEN", "Spinach": "GREEN",
    "Kale": "GREEN", "Capsicum": "GREEN", "Tomato": "GREEN", "Pineapple": "GREEN",
}

COLUMN_ALIASES = {
    "farmer_name":   ["farmer","farmer name","grower","grower name","name","producer",
                      "supplier","full name","farmer_name","outgrower"],
    "crop_type":     ["crop","crop type","produce","product","commodity","item",
                      "produce type","crop_type","variety","description"],
    "net_weight_kg": ["weight","kg","kgs","kilos","kilograms","net weight","gross weight",
                      "weight kg","net_weight","weight_kg","quantity","qty","amount"],
    "origin_county": ["county","region","location","area","origin","district",
                      "sub county","origin_county","county name"],
    "kra_pin":       ["kra pin","kra","pin","tax pin","kra_pin","tax id","tin"],
    "packhouse":     ["packhouse","facility","site","station","pack house","packing station"],
    "notes":         ["notes","remarks","comment","comments","other","additional"],
}

CROP_ALIASES = {
    "avocado": "Avocado", "avo": "Avocado",
    "french beans": "French Beans", "french bean": "French Beans",
    "beans": "French Beans", "fb": "French Beans", "green beans": "French Beans",
    "roses": "Roses", "rose": "Roses",
    "carnations": "Carnations", "carnation": "Carnations",
    "mango": "Mango", "mangoes": "Mango",
    "passion fruit": "Passion Fruit", "passion": "Passion Fruit",
    "macadamia": "Macadamia Nuts", "macadamia nuts": "Macadamia Nuts", "mac": "Macadamia Nuts",
    "tea": "Tea", "coffee": "Coffee",
    "spinach": "Spinach", "kale": "Kale",
    "capsicum": "Capsicum", "pepper": "Capsicum",
    "tomato": "Tomato", "tomatoes": "Tomato",
    "snow peas": "Snow Peas", "snowpeas": "Snow Peas", "peas": "Snow Peas",
    "pineapple": "Pineapple", "maize": "Maize", "corn": "Maize",
}

KRA_PIN_PATTERN = re.compile(r'^[A-Z]\d{9}[A-Z]$')

def load_farmers():
    if os.path.exists(FARMERS_DB):
        with open(FARMERS_DB, "r") as f:
            return json.load(f)
    return {}

def match_farmer(name: str, farmers: dict):
    if not name or not farmers:
        return None
    name_lower = name.strip().lower()
    for fid, fdata in farmers.items():
        if name_lower == fdata["name"].lower():
            return {"id": fid, **fdata}
        parts     = name_lower.split()
        reg_parts = fdata["name"].lower().split()
        if len(parts) >= 2 and any(p in reg_parts for p in parts):
            return {"id": fid, **fdata}
    return None

def normalize_column(col: str) -> str:
    col_clean = str(col).strip().lower()
    col_clean = re.sub(r"[()\[\]/,.]", " ", col_clean)
    col_clean = re.sub(r"[-\s]+", "_", col_clean).strip("_")
    for field, aliases in COLUMN_ALIASES.items():
        alias_clean = [re.sub(r"[-\s]+", "_", re.sub(r"[()\[\]/,.]", " ", a.lower())).strip("_")
                        for a in aliases]
        if col_clean in alias_clean:
            return field
    return col_clean

def _find_header_row(raw_df) -> int:
    best_row, best_score = 0, 0
    for i in range(min(15, len(raw_df))):
        row_vals = [normalize_column(v) for v in raw_df.iloc[i].tolist() if str(v).strip() not in ("", "nan")]
        score = sum(1 for v in row_vals if v in COLUMN_ALIASES.keys())
        if score > best_score:
            best_score, best_row = score, i
    return best_row

def normalize_crop(crop: str) -> str:
    c = str(crop).strip().lower()
    return CROP_ALIASES.get(c, crop.title())

def build_ledger_entries(parsed_records: list, farmers: dict, source: str) -> list:
    entries = []
    for r in parsed_records:
        farmer_name = str(r.get("farmer_name","")).strip()
        crop        = str(r.get("crop_type","")).strip()
        county      = str(r.get("origin_county","")).strip().title()
        kra_pin     = str(r.get("kra_pin","")).strip().upper()
        packhouse   = str(r.get("packhouse","")).strip()
        notes       = str(r.get("notes","")).strip()

        try:
            weight = float(r.get("net_weight_kg", 0) or 0)
        except:
            weight = 0.0

        if not farmer_name or farmer_name in ("nan","None","") or weight <= 0:
            continue

        pin_valid = "✅ Valid" if KRA_PIN_PATTERN.match(kra_pin) else "⚠ Pending"
        hs_code   = HS_CODE_MAP.get(crop, "UNKNOWN")
        fob       = round(weight * 1.5, 2)
        matched   = match_farmer(farmer_name, farmers)

        entries.append({
            "Consignment_ID": f"VP-{datetime.now().strftime('%Y%m%d%H%M%S%f')[:16]}",
            "Timestamp":      datetime.now().strftime("%Y-%m-%d %H:%M"),
            "Farmer_Name":    farmer_name,
            "Farmer_ID":      matched["id"] if matched else "UNREGISTERED",
            "Crop_Type":      crop,
            "KRA_PIN":        kra_pin or "PENDING",
            "PIN_Valid":      pin_valid,
            "HS_Code":        hs_code,
            "Origin_County":  county,
            "Net_Weight_KG":  weight,
            "FOB_Value_USD":  fob,
            "EUDR_Risk":      EUDR_RISK.get(crop, "GREEN"),
            "GPS":            matched.get("gps","") if matched else "",
            "Packhouse":      packhouse,
            "Notes":          notes,
            "Source":         source,
            "_matched":       bool(matched),
        })
    return entries

def get_anthropic_key(input_key: str = "") -> str:
    key = (os.environ.get("ANTHROPIC_API_KEY","") or input_key or "").strip()
    key = key.strip().strip('"').strip("'").strip()
    return key

def get_vision_key(input_key: str = "") -> str:
    key = (os.environ.get("GOOGLE_VISION_API_KEY","") or input_key or "").strip()
    key = key.strip().strip('"').strip("'").strip()
    return key

def validate_anthropic_key(key: str) -> tuple:
    if not key:
        return False, "No API key provided."
    if not key.startswith("sk-ant-"):
        return False, f"Key format looks wrong — should start with 'sk-ant-' (got: {key[:10]}...)"
    if len(key) < 40:
        return False, "Key too short — check for truncation."
    return True, "OK"

def call_claude_parser(raw_content: str, api_key: str) -> list:
    import urllib.request
    import urllib.error

    prompt = f"""You are a data extraction engine for VeriPath Africa, a Kenyan export compliance platform.

Extract ALL farmer/produce records from the raw file content below.

Return a JSON array where each object has these exact keys:
- "farmer_name": string (required)
- "crop_type": string — normalize to one of: Avocado, French Beans, Roses, Carnations, Mango, Passion Fruit, Macadamia Nuts, Tea, Coffee, Spinach, Kale, Capsicum, Tomato, Snow Peas, Pineapple, Maize
- "net_weight_kg": number in kg
- "origin_county": string Kenyan county
- "kra_pin": string or null
- "packhouse": string or null
- "notes": string or null

Rules:
- One object per farmer-crop combination
- If one farmer has multiple crops create multiple objects
- Skip headers totals blank rows
- Return ONLY raw JSON array. No markdown. No explanation. No code fences.

CONTENT:
{raw_content[:7000]}"""

    payload = json.dumps({
        "model":      "claude-sonnet-4-6",
        "max_tokens": 2000,
        "messages":   [{"role": "user", "content": prompt}]
    }).encode("utf-8")

    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=payload,
        headers={
            "content-type":      "application/json",
            "x-api-key":         api_key,
            "anthropic-version": "2023-06-01",
        },
        method="POST"
    )

    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data     = json.loads(resp.read().decode("utf-8"))
            raw_text = data["content"][0]["text"].strip()
            if "```" in raw_text:
                parts    = raw_text.split("```")
                raw_text = parts[1] if len(parts) > 1 else raw_text
                if raw_text.startswith("json"):
                    raw_text = raw_text[4:]
            return json.loads(raw_text.strip())

    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8")
        try:
            err_msg = json.loads(body).get("error",{}).get("message", body)
        except:
            err_msg = body
        st.error(f"API error {e.code}: {err_msg}")
        return []
    except json.JSONDecodeError as e:
        st.error(f"Could not parse AI response as JSON: {e}")
        return []
    except Exception as e:
        st.error(f"Request failed: {e}")
        return []

def rule_based_parse(uploaded_file) -> list:
    name   = uploaded_file.name.lower()
    frames = []
    try:
        if name.endswith(".csv"):
            uploaded_file.seek(0)
            raw = pd.read_csv(uploaded_file, header=None)
            hdr = _find_header_row(raw)
            raw.columns = raw.iloc[hdr]
            frames = [raw.iloc[hdr+1:].reset_index(drop=True)]
        elif name.endswith((".xlsx",".xls")):
            uploaded_file.seek(0)
            xl = pd.ExcelFile(uploaded_file)
            for sheet in xl.sheet_names:
                raw = xl.parse(sheet, header=None)
                if raw.empty:
                    continue
                hdr = _find_header_row(raw)
                df = raw.iloc[hdr+1:].copy()
                df.columns = raw.iloc[hdr]
                if not df.empty:
                    frames.append(df.reset_index(drop=True))
        elif name.endswith(".docx"):
            import docx
            uploaded_file.seek(0)
            doc  = docx.Document(uploaded_file)
            rows = []
            for table in doc.tables:
                headers = [c.text.strip() for c in table.rows[0].cells]
                for row in table.rows[1:]:
                    rows.append({headers[i]: row.cells[i].text.strip()
                                 for i in range(len(headers))})
            if rows:
                frames = [pd.DataFrame(rows)]
    except Exception as e:
        st.error(f"File read error: {e}")
        return []

    records = []
    for df in frames:
        col_map = {col: normalize_column(col) for col in df.columns}
        df      = df.rename(columns=col_map).dropna(how="all")
        for _, row in df.iterrows():
            farmer_name = str(row.get("farmer_name","")).strip()
            crop_raw    = str(row.get("crop_type","")).strip()
            crop        = normalize_crop(crop_raw) if crop_raw not in ("","nan") else ""
            county      = str(row.get("origin_county","")).strip().title()
            kra_pin     = str(row.get("kra_pin","")).strip().upper()
            packhouse   = str(row.get("packhouse","")).strip()
            notes       = str(row.get("notes","")).strip()
            weight_raw  = row.get("net_weight_kg", row.get("weight", 0))
            try:
                weight = float(re.sub(r"[^\d.]","", str(weight_raw)))
            except:
                weight = 0.0
            if not farmer_name or farmer_name in ("nan","None","") or weight <= 0:
                continue
            records.append({
                "farmer_name":   farmer_name,
                "crop_type":     crop or "Unknown",
                "net_weight_kg": weight,
                "origin_county": county if county != "Nan" else "",
                "kra_pin":       kra_pin if kra_pin != "NAN" else "",
                "packhouse":     packhouse if packhouse != "Nan" else "",
                "notes":         notes if notes != "Nan" else "",
            })
    return records

def render_tab_structured(farmers, save_callback):
    st.markdown("#### ⚡ CSV, Excel or Word — instant parse, no API needed")
    st.caption("Works with any column names — auto-mapped to VeriPath schema.")

    uploaded = st.file_uploader(
        "Drop structured file",
        type=["csv","xlsx","xls","docx"],
        key="structured_upload"
    )

    if not uploaded:
        st.markdown("""
        <div style='background:#0d1224;border:2px dashed #1e3a5f;border-radius:12px;
                    padding:32px;text-align:center;color:#64748b'>
            <div style='font-size:1.8rem'>📊</div>
            <div style='margin-top:8px'>CSV · Excel · Word</div>
            <div style='font-size:0.8rem;margin-top:4px'>Any column names — auto-mapped</div>
        </div>
        """, unsafe_allow_html=True)
        return

    st.markdown(f"**File:** `{uploaded.name}` — {uploaded.size/1024:.1f} KB")

    if st.button("⚡ Parse File", use_container_width=True, type="primary", key="parse_structured"):
        with st.spinner("Reading and mapping columns..."):
            raw = rule_based_parse(uploaded)
        if not raw:
            st.error("No valid records found. Check file has farmer name and weight columns.")
            return
        entries = build_ledger_entries(raw, farmers, f"Structured:{uploaded.name}")
        if not entries:
            st.error("Records found but all had missing names or zero weight.")
            return
        st.session_state["ingestion_entries"] = entries
        st.rerun()

def extract_text_from_file(uploaded_file) -> tuple:
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        try:
            import pdfplumber
            uploaded_file.seek(0)
            with pdfplumber.open(uploaded_file) as pdf:
                text = ""
                for page in pdf.pages:
                    pt = page.extract_text()
                    if pt:
                        text += pt + "\n"
                    for table in (page.extract_tables() or []):
                        for row in table:
                            text += " | ".join([str(c).strip() for c in row if c]) + "\n"
            return text.strip(), None
        except ImportError:
            return None, "Run: pip install pdfplumber"
        except Exception as e:
            return None, str(e)
    elif name.endswith(".csv"):
        uploaded_file.seek(0)
        return pd.read_csv(uploaded_file).to_string(index=False), None
    elif name.endswith((".xlsx",".xls")):
        uploaded_file.seek(0)
        xl   = pd.ExcelFile(uploaded_file)
        text = ""
        for sheet in xl.sheet_names:
            text += f"\n=== {sheet} ===\n" + xl.parse(sheet).to_string(index=False)
        return text.strip(), None
    elif name.endswith(".docx"):
        try:
            import docx
            uploaded_file.seek(0)
            doc  = docx.Document(uploaded_file)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            for table in doc.tables:
                for row in table.rows:
                    text += " | ".join([c.text.strip() for c in row.cells]) + "\n"
            return text.strip(), None
        except ImportError:
            return None, "Run: pip install python-docx"
    return None, f"Unsupported: {name.split('.')[-1].upper()}"

def render_tab_ai(farmers, save_callback):
    st.markdown("#### 🧠 AI Parser — messy PDFs and typed documents")
    st.caption("Claude reads any layout. Needs Anthropic API credits at console.anthropic.com")

    env_key = get_anthropic_key()
    if env_key:
        st.success(f"✅ API key loaded from environment ({env_key[:12]}...)")
        api_key = env_key
    else:
        raw_input = st.text_input(
            "Anthropic API Key",
            type="password",
            placeholder="sk-ant-api03-...",
            key="ai_api_key_input"
        )
        api_key = get_anthropic_key(raw_input)

        if raw_input and not api_key.startswith("sk-ant-"):
            st.error(f"Key format wrong — must start with 'sk-ant-' — got: '{api_key[:15]}...'")

    uploaded = st.file_uploader(
        "Drop document",
        type=["pdf","csv","xlsx","xls","docx"],
        key="ai_upload"
    )

    if not uploaded:
        st.markdown("""
        <div style='background:#0d1224;border:2px dashed #1e3a5f;border-radius:12px;
                    padding:32px;text-align:center;color:#64748b'>
            <div style='font-size:1.8rem'>📄</div>
            <div style='margin-top:8px'>PDF · CSV · Excel · Word</div>
            <div style='font-size:0.8rem;margin-top:4px'>Any layout — AI reads it</div>
        </div>
        """, unsafe_allow_html=True)
        return

    st.markdown(f"**File:** `{uploaded.name}` — {uploaded.size/1024:.1f} KB")

    if st.button("🧠 Parse with AI", use_container_width=True, type="primary", key="parse_ai"):
        valid, msg = validate_anthropic_key(api_key)
        if not valid:
            st.error(f"❌ {msg}")
            return

        with st.spinner("Reading file..."):
            raw_text, err = extract_text_from_file(uploaded)
        if err:
            st.error(f"❌ {err}")
            return
        if not raw_text or len(raw_text.strip()) < 20:
            st.error("File appears empty or unreadable.")
            return

        st.success(f"✅ {len(raw_text)} characters extracted")

        with st.spinner("🧠 AI extracting farmer records..."):
            parsed = call_claude_parser(raw_text, api_key)
        if not parsed:
            return

        entries = build_ledger_entries(parsed, farmers, f"AI:{uploaded.name}")
        if not entries:
            st.warning("No valid records extracted.")
            return

        st.session_state["ingestion_entries"] = entries
        st.rerun()

def call_google_vision(image_bytes: bytes, vision_key: str) -> str:
    import urllib.request
    import urllib.error

    b64     = base64.b64encode(image_bytes).decode("utf-8")
    payload = json.dumps({
        "requests": [{
            "image":    {"content": b64},
            "features": [{"type": "DOCUMENT_TEXT_DETECTION", "maxResults": 1}]
        }]
    }).encode("utf-8")

    url = f"https://vision.googleapis.com/v1/images:annotate?key={vision_key}"
    req = urllib.request.Request(
        url, data=payload,
        headers={"Content-Type":"application/json"},
        method="POST"
    )
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            return data["responses"][0].get("fullTextAnnotation",{}).get("text","")
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8")
        st.error(f"Google Vision error {e.code}: {body}")
        return ""
    except Exception as e:
        st.error(f"Vision OCR failed: {e}")
        return ""

def get_image_bytes(uploaded_file) -> tuple:
    name = uploaded_file.name.lower()
    uploaded_file.seek(0)

    if name.endswith((".jpg",".jpeg",".png",".webp",".bmp",".tiff",".tif")):
        return uploaded_file.read(), None

    elif name.endswith(".pdf"):
        try:
            import fitz
            uploaded_file.seek(0)
            doc  = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            page = doc[0]
            pix  = page.get_pixmap(matrix=fitz.Matrix(2,2))
            return pix.tobytes("png"), None
        except ImportError:
            return None, "Run: pip install PyMuPDF"
        except Exception as e:
            return None, str(e)

    return None, f"Unsupported OCR format: {name.split('.')[-1].upper()}"

def render_tab_ocr(farmers, save_callback):
    st.markdown("#### 📷 OCR Scanner — photos and handwritten sheets")
    st.caption("Photograph any handwritten weight log. Google Vision reads it. AI structures it into farmer rows.")

    with st.expander("🔧 Google Vision API Setup"):
        st.markdown("""
        1. Go to **console.cloud.google.com**
        2. Create project: **VeriPath Africa**
        3. Search **Cloud Vision API** → Enable
        4. **APIs & Services → Credentials → Create API Key**
        5. Paste key below
        6. **Free tier: 1,000 images/month** — enough for pilot
        """)

    env_vision = get_vision_key()
    if env_vision:
        st.success(f"✅ Google Vision key loaded from environment ({env_vision[:8]}...)")
        vision_key = env_vision
    else:
        vision_raw = st.text_input(
            "Google Vision API Key",
            type="password",
            placeholder="AIzaSy...",
            key="vision_key_input"
        )
        vision_key = get_vision_key(vision_raw)

    env_claude = get_anthropic_key()
    if env_claude:
        claude_key = env_claude
    else:
        claude_raw = st.text_input(
            "Anthropic API Key (to structure OCR output)",
            type="password",
            placeholder="sk-ant-...",
            key="ocr_claude_input"
        )
        claude_key = get_anthropic_key(claude_raw)

    st.markdown("---")

    uploaded = st.file_uploader(
        "Upload photo or scan",
        type=["jpg","jpeg","png","webp","bmp","tiff","tif","pdf"],
        key="ocr_upload",
        help="Phone photo of handwritten sheet · Scanned PDF · Any image"
    )

    if not uploaded:
        st.markdown("""
        <div style='background:#0d1224;border:2px dashed #1e3a5f;border-radius:12px;
                    padding:32px;text-align:center;color:#64748b'>
            <div style='font-size:1.8rem'>📷</div>
            <div style='margin-top:8px'>Photo · Scan · Handwritten sheet</div>
            <div style='font-size:0.8rem;margin-top:4px'>JPG · PNG · PDF scan · TIFF · WebP</div>
        </div>
        """, unsafe_allow_html=True)
        return

    st.markdown(f"**File:** `{uploaded.name}` — {uploaded.size/1024:.1f} KB")

    name = uploaded.name.lower()
    if not name.endswith(".pdf"):
        uploaded.seek(0)
        st.image(uploaded, caption="Document preview", use_column_width=True)
        uploaded.seek(0)

    st.markdown("---")

    if st.button("📷 Scan & Extract", use_container_width=True, type="primary", key="parse_ocr"):
        if not vision_key:
            st.error("Enter your Google Vision API key.")
            return

        valid, msg = validate_anthropic_key(claude_key)
        if not valid:
            st.error(f"Anthropic key error: {msg}")
            return

        with st.spinner("Preparing image..."):
            img_bytes, err = get_image_bytes(uploaded)
        if err:
            st.error(f"❌ {err}")
            return

        with st.spinner("📖 Google Vision reading text..."):
            raw_text = call_google_vision(img_bytes, vision_key)

        if not raw_text or len(raw_text.strip()) < 10:
            st.error("No text detected. Make sure image is clear and well-lit.")
            return

        st.success(f"✅ OCR complete — {len(raw_text)} characters read")

        with st.expander("📝 Raw OCR output"):
            st.text(raw_text)

        with st.spinner("🧠 AI structuring into farmer rows..."):
            parsed = call_claude_parser(raw_text, claude_key)

        if not parsed:
            return

        entries = build_ledger_entries(parsed, farmers, f"OCR:{uploaded.name}")
        if not entries:
            st.warning("No valid farmer records found in document.")
            return

        st.session_state["ingestion_entries"] = entries
        st.rerun()

def render_preview(save_callback):
    entries = st.session_state.get("ingestion_entries",[])
    if not entries:
        return

    matched   = sum(1 for e in entries if e.get("_matched"))
    unmatched = len(entries) - matched

    st.markdown("---")
    st.markdown("### ✅ Extraction Complete — Preview & Import")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown(f"""<div class='metric-card'>
            <div class='metric-label'>Records Found</div>
            <div class='metric-value'>{len(entries)}</div>
        </div>""", unsafe_allow_html=True)
    with col2:
        tw = sum(e.get("Net_Weight_KG",0) for e in entries)
        st.markdown(f"""<div class='metric-card'>
            <div class='metric-label'>Total Weight KG</div>
            <div class='metric-value'>{tw:,.0f}</div>
        </div>""", unsafe_allow_html=True)
    with col3:
        st.markdown(f"""<div class='metric-card' style='border-color:#16a34a'>
            <div class='metric-label'>✅ Matched</div>
            <div class='metric-value' style='color:#4ade80'>{matched}</div>
        </div>""", unsafe_allow_html=True)
    with col4:
        c = "#fbbf24" if unmatched else "#4ade80"
        st.markdown(f"""<div class='metric-card' style='border-color:{c}'>
            <div class='metric-label'>⚠ Unregistered</div>
            <div class='metric-value' style='color:{c}'>{unmatched}</div>
        </div>""", unsafe_allow_html=True)

    st.markdown("---")

    preview_cols = ["Farmer_Name","Farmer_ID","Crop_Type","Net_Weight_KG",
                    "Origin_County","KRA_PIN","PIN_Valid","HS_Code","EUDR_Risk","Packhouse"]
    preview_df = pd.DataFrame(entries)
    preview_df = preview_df[[c for c in preview_cols if c in preview_df.columns]]

    def color_eudr(val):
        return {"GREEN":"color:#4ade80","AMBER":"color:#fbbf24","RED":"color:#f87171"}.get(val,"")

    st.dataframe(
        preview_df.style.map(color_eudr, subset=["EUDR_Risk"] if "EUDR_Risk" in preview_df.columns else []),
        use_container_width=True,
        height=380
    )

    if unmatched:
        st.warning(f"⚠ {unmatched} farmer(s) not in Outgrower Registry — importing as UNREGISTERED.")

    st.markdown("---")
    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("💾 Import All to Ledger", use_container_width=True, type="primary", key="import_all"):
            if save_callback:
                clean = [{k:v for k,v in e.items() if not k.startswith("_")} for e in entries]
                save_callback(clean)
                st.success(f"✅ {len(entries)} records imported.")
                st.session_state["ingestion_entries"] = []
                st.balloons()
                st.rerun()

    with col2:
        clean_df = pd.DataFrame([{k:v for k,v in e.items() if not k.startswith("_")} for e in entries])
        st.download_button(
            "⬇ Download CSV",
            data=clean_df.to_csv(index=False).encode("utf-8"),
            file_name=f"VeriPath_Parsed_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            mime="text/csv",
            use_container_width=True,
            key="download_parsed"
        )

    with col3:
        if st.button("🗑 Clear", use_container_width=True, key="clear_all"):
            st.session_state["ingestion_entries"] = []
            st.rerun()

def render_data_ingestion_page(save_callback=None):
    st.markdown("# 📥 Data Ingestion")
    st.markdown("<p style='color:#64748b'>Three ways to get farmer data in — structured files, AI documents, or OCR from photos</p>", unsafe_allow_html=True)

    farmers = load_farmers()

    tab1, tab2, tab3 = st.tabs([
        "⚡ Structured Files",
        "🧠 AI Document Parser",
        "📷 OCR Scanner",
    ])

    with tab1:
        render_tab_structured(farmers, save_callback)

    with tab2:
        render_tab_ai(farmers, save_callback)

    with tab3:
        render_tab_ocr(farmers, save_callback)

    render_preview(save_callback)
