import streamlit as st
import pandas as pd
import json
import os
import re
from datetime import datetime

FARMERS_DB = "farmers.json"

HS_CODE_MAP = {
    "Maize": "1005.90", "Coffee": "0901.11", "Tea": "0902.30",
    "Avocado": "0804.40", "French Beans": "0708.20", "Roses": "0603.11",
    "Macadamia Nuts": "0802.62", "Mango": "0804.50",
    "Pineapple": "0804.30", "Passion Fruit": "0810.90",
    "Snow Peas": "0710.21", "Carnations": "0603.12",
    "Spinach": "0709.70", "Kale": "0704.99",
    "Capsicum": "0709.60", "Tomato": "0702.00",
}

EUDR_RISK = {
    "Coffee": "AMBER", "Tea": "AMBER", "Maize": "AMBER",
    "Avocado": "GREEN", "Mango": "GREEN", "Macadamia Nuts": "GREEN",
    "French Beans": "GREEN", "Snow Peas": "GREEN", "Roses": "GREEN",
    "Carnations": "GREEN", "Passion Fruit": "GREEN", "Spinach": "GREEN",
    "Kale": "GREEN", "Capsicum": "GREEN", "Tomato": "GREEN", "Pineapple": "GREEN",
}

# Fuzzy column name aliases — maps any exporter column name to VeriPath field
COLUMN_ALIASES = {
    "farmer_name":    ["farmer","farmer name","grower","grower name","name","producer","supplier","full name","farmer_name"],
    "crop_type":      ["crop","crop type","produce","product","commodity","item","produce type","crop_type","variety"],
    "net_weight_kg":  ["weight","kg","kgs","kilos","kilograms","net weight","gross weight","weight kg","net_weight","weight_kg","quantity","qty","amount"],
    "origin_county":  ["county","region","location","area","origin","district","sub county","origin_county","county name"],
    "kra_pin":        ["kra pin","kra","pin","tax pin","kra_pin","tax id","tin"],
    "packhouse":      ["packhouse","facility","site","station","pack house","packing station"],
    "notes":          ["notes","remarks","comment","comments","other","additional"],
}

CROP_ALIASES = {
    "avocado": "Avocado", "avo": "Avocado",
    "french beans": "French Beans", "french bean": "French Beans", "beans": "French Beans", "fb": "French Beans",
    "roses": "Roses", "rose": "Roses",
    "carnations": "Carnations", "carnation": "Carnations",
    "mango": "Mango", "mangoes": "Mango",
    "passion fruit": "Passion Fruit", "passion": "Passion Fruit",
    "macadamia": "Macadamia Nuts", "macadamia nuts": "Macadamia Nuts", "mac": "Macadamia Nuts",
    "tea": "Tea",
    "coffee": "Coffee",
    "spinach": "Spinach",
    "kale": "Kale",
    "capsicum": "Capsicum", "pepper": "Capsicum",
    "tomato": "Tomato", "tomatoes": "Tomato",
    "snow peas": "Snow Peas", "snowpeas": "Snow Peas", "peas": "Snow Peas",
    "pineapple": "Pineapple",
    "maize": "Maize", "corn": "Maize",
}

KRA_PIN_PATTERN = re.compile(r'^[A-Z]\d{9}[A-Z]$')

def load_farmers():
    if os.path.exists(FARMERS_DB):
        with open(FARMERS_DB, "r") as f:
            return json.load(f)
    return {}

def match_farmer(name: str, farmers: dict):
    if not name or not farmers:
        return None
    name_lower = name.strip().lower()
    for fid, fdata in farmers.items():
        if name_lower == fdata["name"].lower():
            return {"id": fid, **fdata}
        parts     = name_lower.split()
        reg_parts = fdata["name"].lower().split()
        if len(parts) >= 2 and any(p in reg_parts for p in parts):
            return {"id": fid, **fdata}
    return None

def normalize_column(col: str) -> str:
    """Map any column name to a VeriPath field name."""
    col_clean = str(col).strip().lower().replace("-","_").replace(" ","_")
    for field, aliases in COLUMN_ALIASES.items():
        if col_clean in [a.replace(" ","_").replace("-","_") for a in aliases]:
            return field
        if col_clean in aliases:
            return field
    return col_clean

def normalize_crop(crop: str) -> str:
    """Map any crop name to standard VeriPath crop."""
    c = str(crop).strip().lower()
    return CROP_ALIASES.get(c, crop.title())

def extract_file_content(uploaded_file):
    """Extract text for AI parsing."""
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        try:
            import pdfplumber
            uploaded_file.seek(0)
            with pdfplumber.open(uploaded_file) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    for table in (page.extract_tables() or []):
                        for row in table:
                            text += " | ".join([str(c).strip() for c in row if c]) + "\n"
            return text.strip(), None
        except ImportError:
            return None, "pdfplumber not installed — run: pip install pdfplumber"
        except Exception as e:
            return None, f"PDF read error: {e}"

    elif name.endswith(".csv"):
        uploaded_file.seek(0)
        df = pd.read_csv(uploaded_file)
        return df.to_string(index=False), df

    elif name.endswith((".xlsx", ".xls")):
        uploaded_file.seek(0)
        xl   = pd.ExcelFile(uploaded_file)
        text = ""
        for sheet in xl.sheet_names:
            df    = xl.parse(sheet)
            text += f"\n=== Sheet: {sheet} ===\n" + df.to_string(index=False)
        return text.strip(), None

    elif name.endswith(".docx"):
        try:
            import docx
            uploaded_file.seek(0)
            doc  = docx.Document(uploaded_file)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            for table in doc.tables:
                for row in table.rows:
                    text += " | ".join([c.text.strip() for c in row.cells]) + "\n"
            return text.strip(), None
        except ImportError:
            return None, "python-docx not installed"

    return None, f"Unsupported: {name.split('.')[-1].upper()}"

def rule_based_parse(uploaded_file) -> list:
    """
    No API needed — fuzzy column matching for CSV and Excel.
    Works even with messy column names from any exporter.
    """
    name = uploaded_file.name.lower()
    frames = []

    try:
        if name.endswith(".csv"):
            uploaded_file.seek(0)
            frames = [pd.read_csv(uploaded_file)]

        elif name.endswith((".xlsx",".xls")):
            uploaded_file.seek(0)
            xl = pd.ExcelFile(uploaded_file)
            for sheet in xl.sheet_names:
                df = xl.parse(sheet)
                if not df.empty:
                    frames.append(df)
        else:
            return []
    except Exception as e:
        st.error(f"File read error: {e}")
        return []

    records = []
    for df in frames:
        # Normalize column names
        col_map = {col: normalize_column(col) for col in df.columns}
        df = df.rename(columns=col_map)

        # Drop completely empty rows
        df = df.dropna(how="all")

        for _, row in df.iterrows():
            farmer_name = str(row.get("farmer_name","")).strip()
            crop_raw    = str(row.get("crop_type","")).strip()
            crop        = normalize_crop(crop_raw) if crop_raw and crop_raw != "nan" else ""
            county      = str(row.get("origin_county","")).strip().title()
            kra_pin     = str(row.get("kra_pin","")).strip().upper()
            packhouse   = str(row.get("packhouse","")).strip()
            notes       = str(row.get("notes","")).strip()

            # Weight — try to extract number
            weight_raw = row.get("net_weight_kg", row.get("weight", 0))
            try:
                weight = float(re.sub(r"[^\d.]","", str(weight_raw)))
            except:
                weight = 0.0

            if not farmer_name or farmer_name in ("nan","None","") or weight <= 0:
                continue
            if not crop or crop == "Nan":
                crop = "Unknown"

            records.append({
                "farmer_name":    farmer_name,
                "crop_type":      crop,
                "net_weight_kg":  weight,
                "origin_county":  county if county != "Nan" else "",
                "kra_pin":        kra_pin if kra_pin != "NAN" else "",
                "packhouse":      packhouse if packhouse != "Nan" else "",
                "notes":          notes if notes != "Nan" else "",
            })

    return records

def call_claude_parser(raw_content: str, api_key: str) -> list:
    """AI parsing for PDFs and messy files."""
    import urllib.request, urllib.error

    prompt = f"""You are a data extraction engine for VeriPath Africa, a Kenyan export compliance platform.

Extract ALL farmer/produce records from the raw file content below.

Return a JSON array where each object has these exact keys:
- "farmer_name": string — full name of farmer or grower (required)
- "crop_type": string — normalize to one of exactly: Avocado, French Beans, Roses, Carnations, Mango, Passion Fruit, Macadamia Nuts, Tea, Coffee, Spinach, Kale, Capsicum, Tomato, Snow Peas, Pineapple, Maize
- "net_weight_kg": number — weight in kg (convert from other units if needed)
- "origin_county": string — Kenyan county name
- "kra_pin": string or null
- "packhouse": string or null
- "notes": string or null

Rules:
- One object per farmer-crop combination
- Skip header rows, total rows, blank rows
- Return ONLY the raw JSON array. No explanation. No markdown. No code fences.

RAW CONTENT:
{raw_content[:7000]}"""

    payload = json.dumps({
        "model": "claude-sonnet-4-6",
        "max_tokens": 2000,
        "messages": [{"role": "user", "content": prompt}]
    }).encode("utf-8")

    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=payload,
        headers={
            "Content-Type":      "application/json",
            "x-api-key":         api_key,
            "anthropic-version": "2023-06-01",
        },
        method="POST"
    )

    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data     = json.loads(resp.read().decode("utf-8"))
            raw_text = data["content"][0]["text"].strip()
            if "```" in raw_text:
                parts    = raw_text.split("```")
                raw_text = parts[1] if len(parts) > 1 else raw_text
                if raw_text.startswith("json"):
                    raw_text = raw_text[4:]
            return json.loads(raw_text.strip())
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8")
        err  = json.loads(body).get("error",{}).get("message","Unknown error")
        st.error(f"API error: {err}")
        return []
    except json.JSONDecodeError:
        st.error("Could not parse AI response. Try again.")
        return []
    except Exception as e:
        st.error(f"Request failed: {e}")
        return []

def build_ledger_entries(parsed_records: list, farmers: dict, source: str) -> list:
    entries = []
    for r in parsed_records:
        farmer_name = str(r.get("farmer_name","")).strip()
        crop        = str(r.get("crop_type","")).strip()
        county      = str(r.get("origin_county","")).strip().title()
        kra_pin     = str(r.get("kra_pin","")).strip().upper()
        packhouse   = str(r.get("packhouse","")).strip()
        notes       = str(r.get("notes","")).strip()

        try:
            weight = float(r.get("net_weight_kg",0) or 0)
        except:
            weight = 0.0

        if not farmer_name or weight <= 0:
            continue

        pin_valid = "✅ Valid" if KRA_PIN_PATTERN.match(kra_pin) else "⚠ Pending"
        hs_code   = HS_CODE_MAP.get(crop, "UNKNOWN")
        fob       = round(weight * 1.5, 2)
        matched   = match_farmer(farmer_name, farmers)

        entries.append({
            "Consignment_ID": f"VP-{datetime.now().strftime('%Y%m%d%H%M%S%f')[:16]}",
            "Timestamp":      datetime.now().strftime("%Y-%m-%d %H:%M"),
            "Farmer_Name":    farmer_name,
            "Farmer_ID":      matched["id"] if matched else "UNREGISTERED",
            "Crop_Type":      crop,
            "KRA_PIN":        kra_pin or "PENDING",
            "PIN_Valid":      pin_valid,
            "HS_Code":        hs_code,
            "Origin_County":  county,
            "Net_Weight_KG":  weight,
            "FOB_Value_USD":  fob,
            "EUDR_Risk":      EUDR_RISK.get(crop,"GREEN"),
            "GPS":            matched.get("gps","") if matched else "",
            "Packhouse":      packhouse,
            "Notes":          notes,
            "Source":         source,
            "_matched":       bool(matched),
        })
    return entries

def render_smart_parser_page(save_callback=None):
    st.markdown("# 🧠 Smart Data Ingestion")
    st.markdown("<p style='color:#64748b'>Upload any Excel, PDF, CSV or Word — auto-extracts all farmer records</p>", unsafe_allow_html=True)

    farmers = load_farmers()

    # ── Mode selector ──────────────────────────────────────────────────────
    st.markdown("### Parse Mode")
    mode = st.radio(
        "Choose parsing method",
        ["⚡ Fast Parse (CSV/Excel — no API needed)", "🧠 AI Parse (PDF + messy files — needs API credits)"],
        horizontal=True
    )

    api_key = ""
    if "AI Parse" in mode:
        st.markdown("---")
        env_key = os.environ.get("ANTHROPIC_API_KEY","")
        if env_key:
            st.success("✅ ANTHROPIC_API_KEY loaded from environment.")
            api_key = env_key
        else:
            api_key = st.text_input(
                "Anthropic API Key",
                type="password",
                placeholder="sk-ant-...",
                help="Get credits at console.anthropic.com → Billing"
            )
            if not api_key:
                st.markdown("""
                <div style='background:#1a1400;border:1px solid #d97706;border-radius:10px;padding:14px 18px;margin-top:8px'>
                    <div style='color:#fbbf24;font-weight:600'>⚠ API credits needed for AI mode</div>
                    <div style='color:#94a3b8;font-size:0.85rem;margin-top:4px'>
                        Go to <b>console.anthropic.com → Billing</b> and add $5 credits.<br>
                        Or use <b>Fast Parse</b> mode for CSV/Excel files — works right now, no API needed.
                    </div>
                </div>
                """, unsafe_allow_html=True)

    st.markdown("---")

    # ── File upload ────────────────────────────────────────────────────────
    accepted = ["csv","xlsx","xls"] if "Fast" in mode else ["csv","xlsx","xls","pdf","docx"]

    uploaded = st.file_uploader(
        "Drop file here",
        type=accepted,
        help="Fast mode: CSV/Excel only. AI mode: any format including PDF."
    )

    if not uploaded:
        st.markdown(f"""
        <div style='background:#0d1224;border:2px dashed #1e3a5f;border-radius:12px;
                    padding:40px;text-align:center;color:#64748b;margin-top:12px'>
            <div style='font-size:2rem'>📂</div>
            <div style='margin-top:8px'>Drop your packhouse file here</div>
            <div style='font-size:0.8rem;margin-top:6px'>{"CSV · Excel" if "Fast" in mode else "CSV · Excel · PDF · Word"}</div>
        </div>
        """, unsafe_allow_html=True)
        return

    st.markdown(f"**File:** `{uploaded.name}` — {uploaded.size/1024:.1f} KB")
    st.markdown("---")

    if "parsed_entries" not in st.session_state:
        st.session_state.parsed_entries = []

    btn_label = "⚡ Parse File" if "Fast" in mode else "🧠 Parse with AI"

    if st.button(btn_label, use_container_width=True, type="primary"):

        if "Fast" in mode:
            with st.spinner("Reading and mapping columns..."):
                raw_records = rule_based_parse(uploaded)
            if not raw_records:
                st.error("No valid records found. Check file has farmer name and weight columns.")
                return
            entries = build_ledger_entries(raw_records, farmers, f"FastParse:{uploaded.name}")

        else:
            if not api_key:
                st.error("Enter your Anthropic API key to use AI Parse mode.")
                return
            with st.spinner("📖 Reading file..."):
                raw_text, err = extract_file_content(uploaded)
            if err:
                st.error(f"❌ {err}")
                return
            if not raw_text or len(raw_text.strip()) < 20:
                st.error("File appears empty or unreadable.")
                return
            st.success(f"✅ {len(raw_text)} characters extracted from file")
            with st.spinner("🧠 AI extracting farmer records..."):
                parsed = call_claude_parser(raw_text, api_key)
            if not parsed:
                return
            entries = build_ledger_entries(parsed, farmers, f"AIParse:{uploaded.name}")

        if not entries:
            st.warning("No valid records extracted. Check file has farmer names and weights.")
            return

        st.session_state.parsed_entries = entries
        st.rerun()

    # ── Preview ────────────────────────────────────────────────────────────
    if st.session_state.get("parsed_entries"):
        entries   = st.session_state.parsed_entries
        matched   = sum(1 for e in entries if e.get("_matched"))
        unmatched = len(entries) - matched

        st.markdown("### ✅ Extraction Complete")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown(f"""<div class='metric-card'>
                <div class='metric-label'>Records Found</div>
                <div class='metric-value'>{len(entries)}</div>
            </div>""", unsafe_allow_html=True)
        with col2:
            tw = sum(e.get("Net_Weight_KG",0) for e in entries)
            st.markdown(f"""<div class='metric-card'>
                <div class='metric-label'>Total Weight KG</div>
                <div class='metric-value'>{tw:,.0f}</div>
            </div>""", unsafe_allow_html=True)
        with col3:
            st.markdown(f"""<div class='metric-card' style='border-color:#16a34a'>
                <div class='metric-label'>✅ Matched</div>
                <div class='metric-value' style='color:#4ade80'>{matched}</div>
            </div>""", unsafe_allow_html=True)
        with col4:
            c = "#fbbf24" if unmatched else "#4ade80"
            st.markdown(f"""<div class='metric-card' style='border-color:{c}'>
                <div class='metric-label'>⚠ Unregistered</div>
                <div class='metric-value' style='color:{c}'>{unmatched}</div>
            </div>""", unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### Preview — verify before importing")

        preview_cols = ["Farmer_Name","Farmer_ID","Crop_Type","Net_Weight_KG",
                        "Origin_County","KRA_PIN","PIN_Valid","HS_Code","EUDR_Risk","Packhouse"]
        preview_df = pd.DataFrame(entries)
        preview_df = preview_df[[c for c in preview_cols if c in preview_df.columns]]

        def color_eudr(val):
            return {"GREEN":"color:#4ade80","AMBER":"color:#fbbf24","RED":"color:#f87171"}.get(val,"")

        st.dataframe(
            preview_df.style.applymap(color_eudr, subset=["EUDR_Risk"]),
            use_container_width=True, height=380
        )

        if unmatched:
            st.warning(f"⚠ {unmatched} farmer(s) not in Outgrower Registry — will import as UNREGISTERED.")

        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 Import All to Ledger", use_container_width=True, type="primary"):
                if save_callback:
                    clean = [{k:v for k,v in e.items() if not k.startswith("_")} for e in entries]
                    save_callback(clean)
                    st.success(f"✅ {len(entries)} records imported.")
                    st.session_state.parsed_entries = []
                    st.balloons()
                    st.rerun()
                else:
                    st.error("Save callback not connected.")
        with col2:
            if st.button("🗑️ Clear & Re-upload", use_container_width=True):
                st.session_state.parsed_entries = []
                st.rerun()

        st.markdown("---")
        clean_df = pd.DataFrame([{k:v for k,v in e.items() if not k.startswith("_")} for e in entries])
        st.download_button(
            "⬇️ Download Extracted CSV",
            data=clean_df.to_csv(index=False).encode("utf-8"),
            file_name=f"VeriPath_Parsed_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            mime="text/csv",
            use_container_width=True
        )
